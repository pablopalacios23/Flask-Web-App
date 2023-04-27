import os
import json
from flask import Flask, jsonify, render_template, request
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv

load_dotenv('./.flaskenv')

app = Flask(__name__)

@app.route('/', methods=['GET', 'POST'])
def index():
    preguntas = [
        {"tipo": "texto", "pregunta": "¿Cuál es tu nombre?", "nombre": "nombre", "respuesta": ""},
        {"tipo": "seleccion", "pregunta": "¿Eres mayor de edad?", "nombre": "mayor_edad", "opciones": ["Sí", "No"], "respuesta": ""}
    ]

    if request.method == 'POST':
        respuestas = {}
        for pregunta in preguntas:
            if pregunta['tipo'] == 'texto':
                respuesta = request.form[pregunta['nombre']]
                pregunta['respuesta'] = respuesta
                respuestas[pregunta['pregunta']] = respuesta
            elif pregunta['tipo'] == 'seleccion':
                respuesta = request.form[pregunta['nombre']]
                pregunta['respuesta'] = respuesta
                respuestas[pregunta['pregunta']] = respuesta
        with open('respuestas.json', 'w') as f:
            json.dump(respuestas, f)

        # Crear documento Word con respuestas
        document = Document()
        document.add_heading('Respuestas de la encuesta')
        for pregunta in preguntas:
            document.add_heading(pregunta['pregunta'], level=2)
            document.add_paragraph(pregunta['respuesta'])
        document.save('respuestas.docx')

        return 'Formulario enviado correctamente'

    return render_template('index.html', preguntas=preguntas)

if __name__ == '__main__':
    app.run()
